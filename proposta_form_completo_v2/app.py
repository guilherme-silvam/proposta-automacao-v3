from flask import Flask, render_template, request, send_file, flash, redirect, url_for, jsonify
from werkzeug.utils import secure_filename
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Cm
import io, re

app = Flask(__name__)
app.secret_key = "change-this-key"

FIELDS = [
    "ID_HUBSPOT", "NOME_CLIENTE", "NOME_AUTOR", "NUMERO_REVISAO",
    "DATA_EMISSAO", "DATA_VALIDADE", "NOME_EXECUTIVO", "EMAIL_EXECUTIVO",
    "TELEFONE_EXECUTIVO", "QTD_HORA", "VALOR_HORA", "VALOR_TOTAL",
    "CNPJ_CLIENTE", "VALOR_HORA_EXCEDENTE"
]

def replace_in_paragraph(paragraph, mapping: dict):
    if not paragraph.text:
        return
    text = paragraph.text
    changed = False
    for k, v in mapping.items():
        ph = "{{" + k + "}}"
        if ph in text:
            text = text.replace(ph, str(v))
            changed = True
    if changed:
        for r in paragraph.runs:
            r.text = ""
        paragraph.add_run(text)

def replace_in_table(table, mapping: dict):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                replace_in_paragraph(p, mapping)
            for nested in cell.tables:
                replace_in_table(nested, mapping)

def insert_logo_placeholder(doc: Document, logo_path: Path, marker: str = "[LOGO]"):
    def handle_paragraph(p):
        if marker in (p.text or ""):
            for r in p.runs:
                r.text = ""
            run = p.add_run()
            try:
                run.add_picture(str(logo_path), width=Cm(5.0))
            except Exception:
                run.add_picture(str(logo_path))
            return True
        return False

    def handle_table(t):
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if handle_paragraph(p):
                        return True
                for nested in cell.tables:
                    if handle_table(nested):
                        return True
        return False

    for p in doc.paragraphs:
        if handle_paragraph(p):
            return True
    for t in doc.tables:
        if handle_table(t):
            return True
    for sec in doc.sections:
        if sec.header:
            for p in sec.header.paragraphs:
                if handle_paragraph(p):
                    return True
            for t in sec.header.tables:
                if handle_table(t):
                    return True
        if sec.footer:
            for p in sec.footer.paragraphs:
                if handle_paragraph(p):
                    return True
            for t in sec.footer.tables:
                if handle_table(t):
                    return True
    return False

def fill_docx(template_path: Path, data: dict, logo_path: Path | None = None) -> io.BytesIO:
    doc = Document(template_path)

    for p in doc.paragraphs:
        replace_in_paragraph(p, data)
    for t in doc.tables:
        replace_in_table(t, data)

    for sec in doc.sections:
        if sec.header:
            for p in sec.header.paragraphs:
                replace_in_paragraph(p, data)
            for t in sec.header.tables:
                replace_in_table(t, data)
        if sec.footer:
            for p in sec.footer.paragraphs:
                replace_in_paragraph(p, data)
            for t in sec.footer.tables:
                replace_in_table(t, data)

    if logo_path is not None:
        insert_logo_placeholder(doc, logo_path, marker="[LOGO]")

    out_stream = io.BytesIO()
    doc.save(out_stream)
    out_stream.seek(0)
    return out_stream

PLACEHOLDER_RE = re.compile(r"\{\{([A-Z0-9_]+)\}\}")

@app.route("/debug/scan")
def debug_scan():
    tpl = Path("template.docx")
    if not tpl.exists():
        return jsonify({"ok": False, "error": "template.docx não encontrado"}), 404
    doc = Document(tpl)
    found = set()

    def collect_paragraphs(paragraphs):
        for p in paragraphs:
            for m in PLACEHOLDER_RE.finditer(p.text or ""):
                found.add(m.group(1))

    def collect_table(t):
        for r in t.rows:
            for c in r.cells:
                collect_paragraphs(c.paragraphs)
                for nt in c.tables:
                    collect_table(nt)

    collect_paragraphs(doc.paragraphs)
    for t in doc.tables:
        collect_table(t)
    for s in doc.sections:
        if s.header:
            collect_paragraphs(s.header.paragraphs)
            for t in s.header.tables:
                collect_table(t)
        if s.footer:
            collect_paragraphs(s.footer.paragraphs)
            for t in s.footer.tables:
                collect_table(t)

    return jsonify({"ok": True, "placeholders_encontrados": sorted(found)})

@app.route("/", methods=["GET"])
def index():
    return render_template("form.html")

@app.route("/gerar", methods=["POST"])
def gerar():
    data = {k: (request.form.get(k) or "").strip() for k in FIELDS}

    uploads = Path("uploads")
    uploads.mkdir(exist_ok=True)

    template_file = request.files.get("template_file")
    if template_file and template_file.filename:
        tpl_name = secure_filename(template_file.filename)
        template_path = uploads / tpl_name
        template_file.save(template_path)
    else:
        template_path = Path("template.docx")
        if not template_path.exists():
            flash("Envie um template .docx ou coloque 'template.docx' na raiz do projeto.", "error")
            return redirect(url_for("index"))

    logo_file = request.files.get("logo")
    logo_path = None
    if logo_file and logo_file.filename:
        logo_name = secure_filename(logo_file.filename)
        logo_path = uploads / logo_name
        logo_file.save(logo_path)

    output = fill_docx(template_path, data, logo_path)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_name = f"PROPOSTA_{data.get('NOME_CLIENTE','cliente')}_{ts}.docx"

    return send_file(output, as_attachment=True, download_name=out_name,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == "__main__":
    app.run(host="127.0.0.1", port=5000, debug=True)
